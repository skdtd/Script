# coding=UTF-8


import sys
from datetime import datetime
from typing import Union

import docx
import os

import win32com.client as wc
from docx.table import Table, _Row, _Cell
from setuptools.namespaces import flatten

basePath = os.path.abspath('.')
result_file = os.path.join(basePath, 'result')
temp_dir = '.temp'
history_file = '.history'
sep = ' &&& '
docx_list = []
ps = []
if len(sys.argv) != 0:
    ps = sys.argv


def readKeyWordsForFiles(path: str = os.path.abspath('.')) -> list:
    # read single file
    if os.path.isfile(path):
        with open(path, 'r') as f:
            return [i.replace('\n', '').replace('\r', '') for i in f.readlines()]
    # read multiple files
    key_words = []
    keys = [i for i in os.listdir(path) if i.endswith('.key')]
    for i in keys:
        with open(i, 'r') as f:
            key_words = key_words + [i.replace('\n', '').replace('\r', '') for i in f.readlines()]
    return key_words


def getFileSet(rootDir) -> list:
    list_dirs = os.walk(rootDir)
    info = []
    for root, dirs, files in list_dirs:
        for f in files:
            if f.endswith('doc') or f.endswith('docx'):
                info.append(os.path.join(root, f))
    return info


def createTempDir(path: str = os.path.abspath('.')) -> Union[bytes, str]:
    path = os.path.join(path, temp_dir)
    if not os.path.exists(path):
        os.makedirs(path)
    return path


def doc2Docx(index: int, docPath: str, dist: str) -> str:
    if not docPath.endswith('.doc'):
        print('%d - スキップ: %s' % (index + 1, docPath))
        return ''
    print('%d - 処理する: %s' % (index + 1, docPath))
    word = wc.Dispatch("Word.Application")
    doc = word.Documents.Open(docPath)
    tmp = os.path.join(basePath, dist)
    new_file = docPath.replace(basePath, tmp) + 'x'
    parent = os.path.dirname(new_file)
    if not os.path.exists(parent):
        os.makedirs(parent)
    doc.SaveAs(new_file, 12, False, "", True, "", False, False, False, False)  #
    doc.Close()
    word.Quit()
    return new_file


def readDocx(path: str, result: str, *keys):
    with open(result, 'w+') as res:
        print('読み込: %s' % path)
        file = docx.Document(path)
        for t in file.tables:  # type:Table
            for r in t.rows:  # type:_Row
                for c in r.cells:  # type:_Cell
                    text = c.text.replace('\n', '').replace('\r', '')
                    [res.write(sep.join([path, key, text])) for key in flatten(keys) if text.find(key) != -1]
        tmp = ''
        for par in file.paragraphs:
            str1 = tmp + par.text
            text = str1.replace(' ', '').replace('　', '')
            [res.write(sep.join([path, key, text])) for key in flatten(keys) if text.find(key) != -1]
            tmp = par.text


def getFileList(path: str = os.path.abspath('.'), *suffix) -> list:
    list_dirs = os.walk(path)
    fs = []
    for root, dirs, files in list_dirs:
        for f in files:
            for k in suffix:
                if f.endswith(k):
                    fs.append(os.path.join(root, f))
    return fs


def main():
    print('処理開始: %s' % datetime.strftime(datetime.now(), '%Y-%m-%d %H:%M:%S'))
    print('キーワードを読み込む:')
    key_words = readKeyWordsForFiles()
    print('キーワードを読み込む完了: %s' % key_words)
    tmp = createTempDir()
    print('一時フォルダを作成する: %s' % tmp)
    print('統計ファイル中…')
    file_list = getFileList(os.path.abspath('.'), 'doc', 'docx')
    print('統計ファイル完了、全部%d件' % len(file_list))
    ts = datetime.now()
    print('docファイルをdocxファイルにフォーマット開始: %s' % datetime.strftime(ts, '%Y-%m-%d %H:%M:%S'))
    [docx_list.append(doc2Docx(index, file, temp_dir)) for index, file in enumerate(file_list)]
    te = datetime.now()
    print('docファイルをdocxファイルにフォーマット完了: %s' % datetime.strftime(te, '%Y-%m-%d %H:%M:%S'))
    print('所要時間: %s' % str(te - ts))
    ts = datetime.now()
    print('ファイル内容をチェック開始: %s' % datetime.strftime(ts, '%Y-%m-%d %H:%M:%S'))
    [readDocx(file, result_file, key_words) for file in docx_list if not len(file) == 0 and file is not None]
    te = datetime.now()
    print('ファイル内容をチェック完了: %s' % datetime.strftime(te, '%Y-%m-%d %H:%M:%S'))
    print('所要時間: %s' % str(te - ts))
    print('全作業終了')
    print('一時フォルダ: %s' % tmp[1])
    print('結果報告: %s' % result_file)


if __name__ == '__main__':
    main()
