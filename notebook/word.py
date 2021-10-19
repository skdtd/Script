# coding=UTF-8
import argparse
from datetime import datetime

import docx
import os

import win32com.client as wc
from setuptools.namespaces import flatten

cd = os.path.abspath('.')


def timer(text: str):
    def showTime(fn):
        def showText(*keyword):
            start = datetime.now()
            obj = fn(*keyword)
            end = datetime.now()
            print('%s: %s 時間: %s' % (text, datetime.strftime(end, '%Y-%m-%d %H:%M:%S'), str(end - start)))
            return obj

        return showText

    return showTime


class DocHandler:
    suffix = ('doc', 'docx')

    @timer('get file set')
    def getFileSet(self, work_dir: str) -> list:
        fs = []
        for root, dirs, files in os.walk(work_dir):
            for file in files:
                if file.endswith(DocHandler.suffix):
                    fs.append(os.path.join(root, file))
        return fs

    def __init__(self, work_dir=None):
        if work_dir is None:
            self.work_dir = os.path.abspath('.')
        else:
            self.work_dir = work_dir
        self.file_set = self.getFileSet(self.work_dir)


class DocxHandler:

    def __init__(self):
        self.file_set = None

    @staticmethod
    def read(path: str, *param, sep: str = ',') -> list:
        if str == '' or str is None:
            return []
        res_list = []
        tmp = ''
        file = docx.Document(path)
        for t in file.tables:
            for r in t.rows:
                for par in r.cells:
                    str1 = tmp + par.text
                    text = tmp + str1.replace(' ', '').replace('　', '')
                    if text == '' or text is None:
                        continue
                    for key in flatten(param):
                        if text.find(key) != -1:
                            res_list.append(sep.join([path, key, par.text.replace('\n', '').replace('\r', '')]))
                            tmp = ''
                        else:
                            tmp = par.text
        tmp = ''
        for par in file.paragraphs:
            str1 = tmp + par.text
            text = str1.replace(' ', '').replace('　', '')
            if text == '' or text is None:
                continue
            for key in flatten(param):
                if text.find(key) != -1:
                    res_list.append(sep.join([path, key, par.text.replace('\n', '').replace('\r', '')]))
                    tmp = ''
                else:
                    tmp = par.text
        return res_list

    @timer('read docx file info')
    def readSet(self, file_set, *param):
        size = len(file_set)
        result_list = []
        for index, docx in enumerate(file_set):
            lt = self.read(docx, *param)
            result_list = result_list + lt
            print('%d/%d %s' % (index + 1, size, docx))
        return result_list


class FileHandler:

    @staticmethod
    def getKey(file: str) -> list:
        with open(file, 'r') as f:
            return [j for j in (i.replace('\n', '').replace('\r', '') for i in f.readlines()) if len(j) > 0]

    @timer('get key list')
    def readKeys(self, key_set: str) -> list:
        if os.path.isfile(key_set):
            return self.getKey(key_set)
        else:
            key = [i for i in os.listdir(key_set) if i.endswith('.key')]
            key_list = []
            for k in key:
                key_list = key_list + self.getKey(k)
            return key_list

    @timer('create result dir')
    def createResultDir(self, result_dir=None) -> str:
        if result_dir is None:
            for i in os.listdir(cd):
                if i.lower() is 'result':
                    return os.path.join(os.path.abspath('.'), i.lower())
            result_dir = os.path.join(cd, 'result')
            if not os.path.exists(result_dir):
                os.makedirs(result_dir)
            return result_dir
        else:
            return result_dir

    def __init__(self, key_set: str, result_dir=None, work_dir=None):
        # set work dir
        self.work_dir = work_dir
        self.result_dir = self.createResultDir(result_dir)
        self.key_list = self.readKeys(key_set)

    def doc2Docx(self, src: str) -> str:
        if not src.endswith('.doc'):
            return src
        dist = src.replace(self.work_dir, self.result_dir) + 'x'
        if os.path.exists(dist):
            return dist
        word = wc.Dispatch("Word.Application")
        doc = word.Documents.Open(src)
        parent = os.path.dirname(dist)
        if not os.path.exists(parent):
            os.makedirs(parent)
        doc.SaveAs(dist, 12, False, "", True, "", False, False, False, False)
        doc.Close()
        word.Quit()
        return dist

    @timer('doc save as docx')
    def docSet2Docx(self, file_set) -> list:
        size = len(file_set)
        docx_set = []
        for index, doc in enumerate(file_set):
            docx = self.doc2Docx(doc)
            docx_set.append(docx)
            print('%d/%d %s' % (index + 1, size, docx))
        return docx_set

    @timer('write result')
    def writeResult(self, res_list) -> None:
        with open(os.path.join(self.result_dir, 'result'), 'w') as f:
            [f.write(line + '\n') for line in res_list]


if __name__ == '__main__':
    # get param
    parser = argparse.ArgumentParser()
    parser.add_argument("-k", "--keys")
    parser.add_argument("-r", "--result")
    parser.add_argument("-w", "--work")
    args = parser.parse_args()
    keys = args.keys or input('请输入key文件路径或者所在目录:(default: %s)\n' % cd) or cd
    result = args.result or input(
        '请输入临时目录路径:(default: %s)\n' % os.path.join(cd, 'result')) or os.path.join(cd, 'result')
    work = args.work or input('请输入文件所在目录:(default: %s)\n ' % cd) or cd

    fh = FileHandler(keys, result, work)
    dh = DocHandler(args.work)
    dxh = DocxHandler()
    dxh.file_set = fh.docSet2Docx(dh.file_set)
    res = dxh.readSet(dxh.file_set, fh.key_list)
    fh.writeResult(res)
    os.system(r"start notepad.exe %s" % os.path.join(fh.result_dir, 'result'))
